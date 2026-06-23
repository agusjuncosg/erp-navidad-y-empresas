#!/usr/bin/env python3
"""Genera el informe de Auditoría de Diseño Funcional y Operativo del ERP - NAVIDAD Y EMPRESAS"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── PAGE SETUP ────────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ─── HELPERS ───────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def set_para_border_bottom(para, hex_color, size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color.lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(before))
    spacing.set(qn('w:after'), str(after))
    if line:
        spacing.set(qn('w:line'), str(line))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_run(para, text, bold=False, italic=False, size=11, color=None, font="Calibri"):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        r, g, b = hex_to_rgb(color)
        run.font.color.rgb = RGBColor(r, g, b)
    return run

def heading1(text, color="#1F3864"):
    para = doc.add_paragraph()
    set_para_spacing(para, before=300, after=80)
    set_para_border_bottom(para, color, size=12)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.name = "Calibri"
    r, g, b = hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    return para

def heading2(text, color="#2E5DAD"):
    para = doc.add_paragraph()
    set_para_spacing(para, before=200, after=60)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12.5)
    run.font.name = "Calibri"
    r, g, b = hex_to_rgb(color)
    run.font.color.rgb = RGBColor(r, g, b)
    return para

def body(text, size=11):
    para = doc.add_paragraph()
    set_para_spacing(para, before=40, after=60, line=276)
    add_run(para, text, size=size)
    return para

def bullet(text, bold_prefix=None, color_prefix=None, size=10.5):
    para = doc.add_paragraph(style='List Bullet')
    set_para_spacing(para, before=30, after=40, line=264)
    if bold_prefix:
        pfx_run = para.add_run(bold_prefix + " ")
        pfx_run.bold = True
        pfx_run.font.size = Pt(size)
        pfx_run.font.name = "Calibri"
        if color_prefix:
            r, g, b = hex_to_rgb(color_prefix)
            pfx_run.font.color.rgb = RGBColor(r, g, b)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    return para

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        set_para_spacing(p, before=0, after=0)
        p.add_run("").font.size = Pt(4)

# ─── PORTADA ───────────────────────────────────────────────────────────────────
banner = doc.add_paragraph()
banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(banner, before=0, after=100)
r = banner.add_run("NAVIDAD Y EMPRESAS — ERP")
r.bold = True; r.font.size = Pt(22); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(31, 56, 100)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(sub, before=0, after=40)
r2 = sub.add_run("Auditoria de Diseno Funcional y Operativo")
r2.bold = True; r2.font.size = Pt(15); r2.font.name = "Calibri"
r2.font.color.rgb = RGBColor(46, 93, 173)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(meta, before=0, after=200)
r3 = meta.add_run("Informe elaborado: Junio 2026  |  Rol: Consultor ERP Senior / Analista Funcional / UX / Arquitecto de Producto")
r3.italic = True; r3.font.size = Pt(9.5); r3.font.name = "Calibri"
r3.font.color.rgb = RGBColor(120, 120, 120)

sep = doc.add_paragraph()
set_para_border_bottom(sep, "#1F3864", size=16)
set_para_spacing(sep, before=0, after=180)

# ─── RESUMEN EJECUTIVO ─────────────────────────────────────────────────────────
heading1("Resumen Ejecutivo")
body(
    "Este documento es el resultado de una auditoria completa del ERP de Navidad y Empresas, analizado desde la "
    "perspectiva de diseno funcional, flujos operativos, experiencia de usuario y escalabilidad. La revision abarca "
    "los 10 modulos del sistema: Dashboard, Ventas & Leads, Inventario & Combos, Cola de Armado, Logistica, "
    "Caja & Pagos, Compras & Stock, Gastos & Egresos, Gestion de Personal y Configuracion."
)
body(
    "El ERP muestra una arquitectura creativa y un diseno visual solido para la etapa actual. Sin embargo, antes de "
    "salir a produccion real con carga operativa alta, existen fallas de diseno que requieren atencion prioritaria. "
    "Las mas criticas comprometen la integridad del stock, la persistencia de datos y la escalabilidad del sistema."
)
spacer()

# ─── TABLA RESUMEN ─────────────────────────────────────────────────────────────
heading1("Indice de Hallazgos")

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
for cell, txt in zip(hdr, ["Categoria", "Cantidad de hallazgos", "Prioridad maxima"]):
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(255, 255, 255)

rows_data = [
    ("Fallas graves de diseno", "5", "URGENTE - antes de produccion", "FFE0E0"),
    ("Problemas importantes de usabilidad", "8", "ALTA - afecta operacion diaria", "FFF0E0"),
    ("Mejoras recomendadas", "7", "MEDIA - impacto en eficiencia", "FFFDE0"),
    ("Aspectos bien resueltos", "8", "-", "E8F5E9"),
    ("Reorganizaciones sugeridas", "5", "ALTA - antes de temporada alta", "E3F0FF"),
]
for cat, qty, prio, bg in rows_data:
    row = table.add_row().cells
    for cell, txt in zip(row, [cat, qty, prio]):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if txt == qty else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(txt)
        run.font.size = Pt(10); run.font.name = "Calibri"

widths = [Cm(7.5), Cm(4.2), Cm(5.5)]
for row in table.rows:
    for cell, w in zip(row.cells, widths):
        cell.width = w

spacer(2)

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 1: FALLAS GRAVES
# ══════════════════════════════════════════════════════════════════════════════
heading1("FALLAS GRAVES DE DISENO", color="#C0392B")
body("Problemas que pueden generar perdida de datos, inconsistencias operativas o colapso del sistema a escala. Deben resolverse antes de salir a produccion real.")
spacer()

# FG-01
heading2("FG-01 - El ERP corre 100% en localStorage del navegador", color="#C0392B")
body("Que observo:")
bullet("Toda la persistencia del sistema vive en localStorage del browser (clave: christmas_erp_data_v2).")
bullet("No hay servidor, no hay base de datos real, no hay API.")
bullet("El limite de localStorage es de 5 a 10 MB segun el browser.")
spacer()
body("Por que es un problema:")
bullet("Con 1.000 empresas, 80.000 cajas, miles de movimientos de stock y ordenes de compra, el volumen de datos superara el limite fisico del storage.")
bullet("Si el usuario limpia el historial del navegador, borra cookies o usa un browser diferente, PIERDE TODOS LOS DATOS DE LA EMPRESA.")
bullet("No hay multi-usuario real: cada persona que abra el ERP en su computadora ve una base de datos completamente distinta y desincronizada.")
bullet("No hay respaldo automatico. El boton 'Restaurar Datos de Fabrica' borra todo sin posibilidad de recuperacion.")
bullet("No hay auditoria centralizada ni logs persistentes que sobrevivan entre sesiones.")
spacer()
body("Impacto: CRITICO. Un solo browser corrupto o limpiado destruye toda la operacion.")
body("Como resolverlo: Migrar hacia un backend real (API + base de datos). Como paso intermedio, implementar exportacion e importacion JSON manual como mecanismo de backup hasta tener el backend.")
body("Prioridad: URGENTE - es la base arquitectonica de todo el sistema.", size=10)
spacer()

# FG-02
heading2("FG-02 - El stock NO se reserva al confirmar una venta", color="#C0392B")
body("Que observo:")
bullet("El descuento fisico de stock ocurre unicamente cuando se presiona 'Iniciar Armado' en el modulo de Cola de Armado (assembly.js, linea 178: prod.stock -= qtyNeeded).")
bullet("Entre el momento en que una venta se confirma y el momento en que se inicia el armado, el stock NO esta formalmente reservado.")
bullet("El modulo de inventario calcula visualmente el 'stock comprometido' sumando las ordenes confirmadas y en produccion, pero esto es solo visual: no bloquea operaciones.")
spacer()
body("Por que es un problema:")
bullet("Si hay dos ventas confirmadas que necesitan el mismo insumo y ambas tienen stock disponible al momento de confirmacion, cuando el deposito inicia el segundo armado puede dejar stock en negativo.")
bullet("Una persona puede confirmar una venta que en la pantalla muestra 'stock disponible', pero ese stock ya estaba visualmente comprometido por otra venta y el sistema permite igualmente proceder.")
bullet("El riesgo escala dramaticamente con multiples operadores trabajando en simultaneo.")
spacer()
body("Impacto: Posibles armados con faltantes fisicos de producto, incumplimiento de entregas.")
body("Como resolverlo: Al confirmar una venta, reducir inmediatamente un campo 'stock_reservado' en el producto. El stock disponible real = stock_fisico - stock_reservado. Liberar la reserva solo si se cancela la orden.")
body("Prioridad: URGENTE.", size=10)
spacer()

# FG-03
heading2("FG-03 - No existe trazabilidad entre compras y stock", color="#C0392B")
body("Que observo:")
bullet("Al registrar una factura de compra, el sistema suma las cantidades directamente al campo stock del producto (prod.stock +=).")
bullet("No hay numero de lote, fecha de vencimiento, ni referencia a la factura de origen asociada a las unidades en stock.")
bullet("Si mañana necesitas saber de que factura vienen determinadas unidades de un producto, no hay forma de saberlo desde el sistema.")
spacer()
body("Por que es un problema:")
bullet("Imposible hacer auditoria de stock: que compre, cuando, a que precio y cuanto queda de esa compra.")
bullet("Si un proveedor entrega producto en mal estado, no podes identificar que pedidos usan ese lote.")
bullet("El costo del producto no se actualiza automaticamente al registrar una compra a diferente precio, lo que genera costos incorrectos en los combos.")
spacer()
body("Impacto: Perdida de trazabilidad de inventario. Costos desactualizados. Imposibilidad de auditoria.")
body("Como resolverlo: Agregar un modelo de 'movimientos de stock' con referencia a la compra de origen. El stock actual = suma de todos los movimientos positivos y negativos.")
body("Prioridad: URGENTE.", size=10)
spacer()

# FG-04
heading2("FG-04 - El modulo de Armado mezcla tres areas de responsabilidad", color="#C0392B")
body("Que observo:")
bullet("El modulo 'Cola de Armado' contiene en tres pestanas: (1) Cola de Armado y Produccion, (2) Control de Asistencia Diario, y (3) Liquidacion de Sueldos.")
bullet("Estas tres areas son conceptualmente independientes: produccion, RRHH y administracion de sueldos.")
spacer()
body("Por que es un problema:")
bullet("Un jefe de produccion necesita ver la cola de armado durante todo el dia. Un administrativo necesita ver la liquidacion de sueldos. Estas personas no deberian compartir la misma pantalla.")
bullet("A escala, la pestana de Asistencia puede tener cientos de registros diarios que nada tienen que ver con ver que pedidos armar.")
bullet("La Liquidacion de Sueldos no tiene ninguna relacion funcional con la Cola de Armado. El lugar correcto es dentro de 'Gestion de Personal'.")
spacer()
body("Impacto: Confusion operativa. Dificultad de asignar el modulo al rol correcto.")
body("Como resolverlo: Mover 'Asistencia' y 'Liquidacion' al modulo de Gestion de Personal. 'Cola de Armado' debe ser exclusivamente produccion.")
body("Prioridad: ALTA.", size=10)
spacer()

# FG-05
heading2("FG-05 - Los roles no tienen control de acceso real", color="#C0392B")
body("Que observo:")
bullet("El selector de rol ('Administrador', 'Vendedor Comercial', 'Jefe de Produccion', 'Coordinador de Envios') es un dropdown en el sidebar que cualquier usuario puede cambiar sin restriccion.")
bullet("No hay autenticacion de usuarios. No hay sesiones. No hay permisos reales que bloqueen acciones.")
bullet("Un operario del deposito puede acceder a toda la informacion financiera, margenes, costos y datos de clientes con un clic.")
spacer()
body("Por que es un problema:")
bullet("No hay forma de saber quien hizo que: el log de acciones registra el rol seleccionado en el momento, que cualquiera puede cambiar.")
bullet("En una empresa real con varios empleados, esto es inaceptable desde el punto de vista de seguridad y confidencialidad.")
spacer()
body("Impacto: Seguridad comprometida. Confidencialidad de datos nula.")
body("Como resolverlo: Implementar autenticacion real con login/password. Los permisos deben controlarse en el backend, no en el frontend.")
body("Prioridad: ALTA antes de dar acceso a empleados.", size=10)
spacer(2)

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 2: PROBLEMAS IMPORTANTES
# ══════════════════════════════════════════════════════════════════════════════
heading1("PROBLEMAS IMPORTANTES DE USABILIDAD", color="#E67E22")
body("Problemas que no comprometen la integridad de datos pero generan ineficiencia operativa, errores humanos frecuentes o dificultad de uso en condiciones de alta carga.")
spacer()

issues = [
    ("PI-01 - El CRM (Leads) y la gestion de Pedidos estan en el mismo modulo",
     [
         "El modulo 'Ventas & Leads' tiene dos pestanas: 'Leads y Prospectos' (tablero kanban) y 'Presupuestos y Ventas' (tabla de ordenes).",
         "Son dos universos operativos completamente distintos: el CRM es pre-venta, los pedidos son post-venta.",
         "A escala, con cientos de leads activos y miles de ordenes, la misma persona no puede gestionar ambas cosas desde la misma pantalla.",
         "El boton '+ Nueva Cotizacion' aparece en el encabezado aunque estes en la pestana de Leads, generando confusion.",
     ],
     "Separar en dos modulos del sidebar: 'CRM / Leads' y 'Pedidos / Ventas'. Esto permite asignar distintos roles a cada modulo.",
     "ALTA"),

    ("PI-02 - Las notas de un lead son un unico campo de texto (no un historial)",
     [
         "El campo 'notas' de un lead es un textarea que se sobreescribe cada vez que se presiona 'Nota'.",
         "Si hay 5 llamadas al mismo cliente en distintas fechas, solo se conserva la ultima version del texto.",
         "No hay fecha, hora ni usuario asociado a cada interaccion.",
     ],
     "Convertir las notas en un historial de interacciones: lista de entradas con fecha, hora, usuario y texto. Nunca sobreescribir, siempre agregar.",
     "ALTA"),

    ("PI-03 - Los recordatorios no tienen vista centralizada",
     [
         "Los recordatorios de leads estan ocultos dentro de cada tarjeta kanban. No hay ninguna vista de 'Alertas del dia' ni 'Recordatorios pendientes'.",
         "El Dashboard muestra alertas de riesgo operativo, pero no muestra recordatorios comerciales del dia.",
         "En temporada alta con 50+ leads activos, es imposible saber cuales tienen recordatorios pendientes sin revisar cada tarjeta.",
     ],
     "Agregar un bloque en el Dashboard: 'Recordatorios de Hoy' con el lead, el texto del recordatorio y un acceso directo.",
     "ALTA"),

    ("PI-04 - El modulo de Logistica tiene 5 pestanas que fragmentan la informacion",
     [
         "Las 5 pestanas son: Cordoba Capital (Lista), Otras Provincias (Lista), Calendario Cordoba, Calendario Provincias, Ventas con Datos Pendientes.",
         "Para tener una vision completa de todas las entregas del dia siguiente, hay que navegar entre al menos 4 pantallas.",
         "No hay buscador ni filtro de texto en ninguna de las listas de logistica.",
         "La separacion en calendarios separados por zona obliga a revisar dos calendarios para ver el panorama completo de la semana.",
     ],
     "Unificar en 3 pestanas: 'Lista de Entregas' (con filtro de zona), 'Calendario' (con selector de zona), 'Datos Pendientes'. Agregar buscador por cliente o pedido.",
     "ALTA"),

    ("PI-05 - Los combos no muestran alerta de stock insuficiente para armarlos",
     [
         "En la vista de combos del modulo de Inventario, se muestra el costo de materiales y el margen, pero no si hay stock fisico suficiente para armar una unidad del combo.",
         "Se puede cotizar un combo aunque ninguno de sus ingredientes tenga stock.",
     ],
     "En la tarjeta de cada combo, mostrar el numero de unidades armables con el stock actual (minimo de todos sus ingredientes). Alerta en rojo si es 0.",
     "MEDIA"),

    ("PI-06 - El nombre del modulo 'Compras & Stock' es incorrecto",
     [
         "El modulo de Compras registra facturas de proveedores e ingresa mercaderia, pero el stock se gestiona desde el modulo 'Inventario & Combos'.",
         "Un nuevo usuario que busca 'ajustar stock' ira a 'Compras & Stock' y no encontrara lo que busca.",
         "La funcion de ajuste rapido de stock (+ / -) esta en Inventario, no en Compras.",
     ],
     "Renombrar el modulo a 'Compras' o 'Registro de Facturas'. Centralizar todo el stock en 'Inventario'.",
     "BAJA"),

    ("PI-07 - No hay gestion de devoluciones ni pedidos parciales",
     [
         "El flujo de estados es lineal: Confirmado - En Produccion - Terminado/Listo - Despachado - Entregado.",
         "No existe ningun estado o flujo para: cliente que rechaza parte de un pedido, pedido con productos faltantes, entrega parcial, devolucion.",
         "Si un cliente recibe 80 cajas de las 100 pedidas, no hay forma de registrar ese evento sin cancelar toda la orden.",
     ],
     "Agregar opcion de 'Entrega Parcial' que registre la cantidad efectivamente entregada vs. la pedida, y genere automaticamente un saldo pendiente o nota de credito.",
     "MEDIA"),

    ("PI-08 - El estado 'Terminado / Listo' tiene nomenclatura ambigua",
     [
         "El estado 'Terminado / Listo' sugiere dos posibles interpretaciones: terminado de armar, listo para despachar, o ambas.",
         "La barra '/' es un anti-patron en estados de un sistema: los estados deben ser unicos e inequivocos.",
         "En filtros, reportes y busquedas, este estado con caracteres especiales puede generar problemas tecnicos.",
     ],
     "Renombrar a 'Listo para Despacho'. Es inequivoco y describe exactamente la accion siguiente.",
     "BAJA"),
]

for title, obs_list, solucion, prioridad in issues:
    heading2(title, color="#E67E22")
    body("Observaciones:")
    for obs in obs_list:
        bullet(obs)
    body("Como resolverlo: " + solucion)
    body("Prioridad: " + prioridad + ".", size=10)
    spacer()

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 3: MEJORAS RECOMENDADAS
# ══════════════════════════════════════════════════════════════════════════════
heading1("MEJORAS RECOMENDADAS", color="#B7950B")
body("Oportunidades para reducir carga operativa, eliminar errores humanos y aprovechar mejor la informacion ya existente en el sistema.")
spacer()

mejoras = [
    ("MR-01 - No hay exportacion de datos operativos",
     "No existe ningun boton para exportar a Excel o CSV: el historial de ventas, el listado de stock, las compras, los gastos, los movimientos de caja. "
     "El unico documento que se puede imprimir es la Hoja de Armado y el Remito.",
     "Agregar exportacion a Excel/CSV en: Ventas (listado de ordenes), Inventario (stock actual), Compras (historial de facturas), Caja (movimientos de cobros), Gastos. Es una funcionalidad de alta demanda y bajo costo de implementacion."),

    ("MR-02 - El costo de un producto no se actualiza al registrar una compra a distinto precio",
     "Cuando se registra una factura de compra, el stock del producto aumenta pero el campo 'cost' del producto en el catalogo NO cambia. "
     "Si el mes pasado compre sidra a $1.200 y ahora la compre a $1.600, el sistema sigue calculando el costo del combo con $1.200.",
     "Al registrar una compra, preguntar: '¿Actualizar el costo del producto en el catalogo con el precio de esta factura?' Esto mantiene costos y margenes actualizados sin ir a editar cada producto manualmente."),

    ("MR-03 - Los proveedores no son entidades del sistema",
     "En el modulo de Compras, el proveedor se ingresa como texto libre en cada factura. No hay un ABM de proveedores. "
     "No se puede ver el historial de compras a un proveedor especifico, ni sus condiciones comerciales, ni sus datos de contacto.",
     "Crear un ABM de Proveedores con nombre, CUIT, telefono, condiciones de pago y email. Vincular cada compra a un proveedor del ABM para habilitar busquedas y reportes por proveedor."),

    ("MR-04 - No hay vista de 'Flujo de Caja' o resultado del periodo",
     "El Dashboard muestra KPIs de ventas e indicadores de rentabilidad, pero no hay ninguna vista que muestre: ingresos cobrados en el periodo, egresos del periodo (gastos + compras + sueldos), y resultado neto. "
     "El dueno necesita saber cuanto entro y cuanto salio esta semana/mes.",
     "Agregar una pestana de 'Resultado del Periodo' en el Dashboard o en Caja & Pagos con: Total cobrado, Total de compras, Total de gastos, Total de sueldos liquidados, Resultado neto."),

    ("MR-05 - Los IDs de pedidos son poco legibles operativamente",
     "Los IDs de pedidos son generados como 'ord_' + timestamp (ej: ord_1701234567890). En pantalla se muestra solo el sufijo numerico. "
     "En una operacion real, los operarios y clientes usan un numero de pedido para comunicarse. Un numero de 13 digitos es imposible de memorizar o dictar.",
     "Implementar numeracion correlativa y legible: P-2025-001, P-2025-002, etc. El ID interno puede seguir siendo el timestamp, pero el numero visible debe ser humano."),

    ("MR-06 - No hay busqueda global en el ERP",
     "Si necesitas encontrar una empresa cliente rapidamente, tenes que navegar a Ventas, ir a la pestana correcta, y revisar la lista. "
     "No hay una barra de busqueda que permita buscar 'Holcim' y llegar directamente a su lead o su pedido.",
     "Agregar un buscador global en el header que busque por nombre de empresa en leads, pedidos, compras y gastos simultaneamente."),

    ("MR-07 - El modulo de Empleados muestra 'PRECIOS NETOS SIN IVA' en el encabezado",
     "El modulo de Gestion de Personal tiene en el titulo: 'PRECIOS NETOS SIN IVA'. Los sueldos de los empleados no son 'precios'. "
     "Es un copy-paste del encabezado de otros modulos que quedo sin corregir.",
     "Ajustar el texto del encabezado. Es un detalle menor pero genera mala impresion de prolijidad del sistema."),
]

for title, obs, solucion in mejoras:
    heading2(title, color="#B7950B")
    body("Observacion: " + obs)
    body("Como resolverlo: " + solucion)
    spacer()

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 4: ASPECTOS BIEN RESUELTOS
# ══════════════════════════════════════════════════════════════════════════════
heading1("ASPECTOS BIEN RESUELTOS", color="#1E8449")
body("Decisiones de diseno correctas que deben preservarse y tomarse como referencia para el desarrollo futuro.")
spacer()

bien = [
    ("Sistema de Semaforo por Pedido",
     "El semaforo de colores (verde/amarillo/rojo) en cada pedido es una solucion muy inteligente para gestion de riesgo operativo. "
     "En un vistazo, el operador sabe que pedidos tienen problemas: falta CUIT, falta sena, stock negativo. Es simple, efectivo y dificil de ignorar."),
    ("Disenador de Combos con Analisis de Costo y Margen en Tiempo Real",
     "El modal de creacion de combos muestra en tiempo real el costo de materiales, el precio de venta, el margen en pesos y el margen porcentual mientras se arma el combo. "
     "Esto es exactamente lo que necesita un negocio para no vender a perdida. Esta bien pensado y bien ejecutado."),
    ("Cola de Armado con Ordenamiento Inteligente",
     "La cola de armado ordena automaticamente los pedidos por zona (primero provincias, luego Cordoba) y dentro de cada zona por fecha de entrega mas proxima. "
     "Esto refleja la realidad operativa de la empresa y reduce la carga cognitiva del jefe de deposito."),
    ("Vencimiento Automatico de Presupuestos",
     "El sistema detecta automaticamente presupuestos con mas de 10 dias sin confirmar y los marca como vencidos con alerta en el Dashboard. "
     "Es un control comercial valioso que evita vender a precios desactualizados en un contexto inflacionario."),
    ("Calendarios de Logistica Separados por Zona",
     "Los calendarios visuales separados entre Cordoba Capital y Otras Provincias tienen sentido operativo real: son rutas de distribucion distintas con logisticas distintas. "
     "La visualizacion por mes con puntos de color por estado es clara y funcional."),
    ("Hoja de Armado Imprimible",
     "La Hoja de Armado imprimible con detalle completo del combo (ingredientes, cantidades, destino) es una funcionalidad operativa critica bien implementada. "
     "El diseno apaisado es correcto para uso en deposito."),
    ("Configuracion de Tarifas de Logistica",
     "El modulo de Configuracion permite definir tarifas base y por pallet para distintas zonas (CABA, AMBA, Cordoba Capital). "
     "Esto hace que el calculo logistico sea configurable sin tocar el codigo, lo cual es correcto."),
    ("Diseno Visual y Navegacion",
     "La paleta de colores, la consistencia visual entre modulos, el uso de badges de estado y el diseno general del sidebar son solidos. "
     "El sistema es visualmente legible y no sobrecargado. El 'Panel de Guerra' como nombre del Dashboard es un buen toque para una operacion estacional de alta intensidad."),
]

for title, descripcion in bien:
    heading2("OK  " + title, color="#1E8449")
    body(descripcion)
    spacer()

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 5: REORGANIZACIONES
# ══════════════════════════════════════════════════════════════════════════════
heading1("REORGANIZACIONES SUGERIDAS", color="#2471A3")
body("Cambios de estructura del ERP que mejorarian la navegacion, la asignacion de roles y la escalabilidad sin cambiar la logica de negocio.")
spacer()

reorgs = [
    ("RS-01 - Separar CRM/Leads de Ventas/Pedidos en el sidebar",
     "Actualmente: Un solo modulo 'Ventas & Leads' con dos pestanas.",
     "Propuesta: Dos modulos distintos en el sidebar: 'CRM / Leads' y 'Pedidos / Ventas'. "
     "Esto permite que el vendedor comercial vea su modulo de CRM y el administrativo su modulo de pedidos sin mezclarlos.",
     "Impacto: Navegacion mas clara, mejor asignacion por rol, mas escalable con carga de datos alta.",
     "Modulos afectados: Sidebar, sales.js, app.js"),

    ("RS-02 - Mover Asistencia y Liquidacion al modulo de Gestion de Personal",
     "Actualmente: Cola de Armado tiene 3 pestanas: Produccion, Asistencia, Liquidacion.",
     "Propuesta: Cola de Armado solo tiene produccion. Asistencia y Liquidacion se mueven a Gestion de Personal como pestanas adicionales.",
     "Impacto: El jefe de deposito ve solo la cola. El administrativo gestiona personal desde un solo lugar. Logica de roles mas clara.",
     "Modulos afectados: assembly.js, employees.js, sidebar"),

    ("RS-03 - Unificar Gastos dentro de Caja & Pagos",
     "Actualmente: 'Caja & Pagos' (cobros de ventas) y 'Gastos & Egresos' (egresos) son dos modulos separados.",
     "Propuesta: Unificar en un modulo 'Finanzas & Caja' con tres pestanas: Control de Cobros, Gastos & Egresos, y Resultado del Periodo.",
     "Impacto: El dueno tiene en un solo modulo toda la informacion financiera: entradas, salidas y resultado. Menos clics para el cierre del dia.",
     "Modulos afectados: payments.js, expenses.js, sidebar"),

    ("RS-04 - Simplificar las pestanas del modulo de Logistica",
     "Actualmente: 5 pestanas (Cordoba Lista, Provincias Lista, Calendario Cordoba, Calendario Provincias, Datos Pendientes).",
     "Propuesta: 3 pestanas: 'Lista de Entregas' (con filtro de zona: todas/Cordoba/Provincias), 'Calendario' (con selector de zona), 'Datos Pendientes'.",
     "Impacto: Menos clics para ver el panorama completo. Un solo calendario con toggle es mas eficiente que dos calendarios separados.",
     "Modulos afectados: logistics.js"),

    ("RS-05 - Estructura de sidebar propuesta para la version escalable",
     "Sidebar actual (10 items): Dashboard, Ventas & Leads, Inventario & Combos, Cola de Armado, Logistica, Caja & Pagos, Compras & Stock, Gastos & Egresos, Gestion de Personal, Configuracion.",
     "Sidebar propuesto (9 items con mejor organizacion): Panel de Guerra, CRM / Leads [SEPARADO], Pedidos / Ventas [SEPARADO], Inventario & Combos, Cola de Armado [SOLO PRODUCCION], Logistica [SIMPLIFICADA], Finanzas & Caja [UNIFICADO], Compras [RENOMBRADO], Personal & Sueldos [AMPLIADO], Configuracion.",
     "Impacto: Flujo mas natural Lead - Pedido - Produccion - Logistica - Cobro. Cada modulo tiene una responsabilidad clara.",
     "Modulos afectados: app.js, sidebar, todos los modulos afectados por reorganizaciones RS-01 a RS-04"),
]

for title, actual, propuesta, impacto, afectados in reorgs:
    heading2(title, color="#2471A3")
    body("Situacion actual: " + actual)
    body("Propuesta: " + propuesta)
    body("Impacto: " + impacto)
    body("Modulos afectados: " + afectados + ".", size=10)
    spacer()

spacer()

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 6: ESCALABILIDAD
# ══════════════════════════════════════════════════════════════════════════════
heading1("Evaluacion de Escalabilidad Operativa")
body("Analisis de que sucederia si el sistema opera a escala maxima: 1.000+ empresas, 80.000+ cajas, multiples operadores simultaneos.")
spacer()

heading2("Problemas que hoy no se notan pero que escalaran mal:", color="#6C3483")
bullet("localStorage con miles de ordenes, leads, movimientos de stock y registros de asistencia superara el limite de 5-10 MB. El sistema fallara silenciosamente o perdera datos.", bold_prefix="CRITICO:", color_prefix="#C0392B")
bullet("Las listas de pedidos y leads no tienen paginacion. Con 500 ordenes activas, renderizar toda la tabla en el DOM generara una interfaz lenta e inutilizable.", bold_prefix="CRITICO:", color_prefix="#C0392B")
bullet("El buscador de stock en Inventario hace un filter() sobre todo el catalogo en cada tecla. Con 500+ productos, esto sera perceptiblemente lento.", bold_prefix="IMPORTANTE:", color_prefix="#E67E22")
bullet("Los KPIs del Dashboard recalculan en tiempo real sumando todos los pedidos, todas las compras, todos los gastos. Sin cache, esto sera el cuello de botella de performance principal.", bold_prefix="IMPORTANTE:", color_prefix="#E67E22")
bullet("El modulo de Logistica filtra y renderiza todos los pedidos activos cada vez que se cambia de pestana. A 1.000 pedidos activos, esto genera re-renders lentos.", bold_prefix="MEDIO:", color_prefix="#B7950B")
bullet("Los IDs generados con Math.random() pueden colisionar. A escala, la probabilidad aumenta.", bold_prefix="BAJO:", color_prefix="#1E8449")
spacer()

heading2("Recomendaciones especificas para escalar:", color="#6C3483")
bullet("Implementar un backend real (Node.js/Python + PostgreSQL o similar) con API REST antes de superar los 100 pedidos activos simultaneos.")
bullet("Agregar paginacion en todas las tablas (20 a 50 registros por pagina).")
bullet("Implementar busqueda server-side para listas grandes.")
bullet("Agregar indices y cache de KPIs en el backend en lugar de recalcular en el browser.")
spacer(2)

# ══════════════════════════════════════════════════════════════════════════════
# SECCION 7: TABLA DE PRIORIDADES
# ══════════════════════════════════════════════════════════════════════════════
heading1("Tabla de Prioridades de Implementacion")

priority_table = doc.add_table(rows=1, cols=4)
priority_table.style = 'Table Grid'
hdr_cells = priority_table.rows[0].cells
for cell, txt in zip(hdr_cells, ["Codigo", "Hallazgo", "Categoria", "Prioridad"]):
    set_cell_bg(cell, "1F3864")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(9.5); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(255, 255, 255)

priority_data = [
    ("FG-01", "Migrar de localStorage a backend real", "Falla Grave", "P0 - URGENTE", "FFE0E0"),
    ("FG-02", "Reserva de stock al confirmar venta", "Falla Grave", "P0 - URGENTE", "FFE0E0"),
    ("FG-03", "Trazabilidad compra - stock", "Falla Grave", "P0 - URGENTE", "FFE0E0"),
    ("FG-05", "Autenticacion y control de acceso real", "Falla Grave", "P0 - URGENTE", "FFE0E0"),
    ("FG-04", "Separar responsabilidades del modulo Armado", "Falla Grave", "P1 - ALTA", "FFE0E0"),
    ("RS-01", "Separar CRM de Pedidos en sidebar", "Reorganizacion", "P1 - ALTA", "E3F0FF"),
    ("RS-02", "Mover Asistencia/Liquidacion a Personal", "Reorganizacion", "P1 - ALTA", "E3F0FF"),
    ("PI-01", "Separar CRM de Ventas en modulo propio", "Usabilidad", "P1 - ALTA", "FFF0E0"),
    ("PI-02", "Historial de interacciones en leads", "Usabilidad", "P1 - ALTA", "FFF0E0"),
    ("PI-03", "Recordatorios del dia en Dashboard", "Usabilidad", "P1 - ALTA", "FFF0E0"),
    ("PI-04", "Simplificar pestanas de Logistica", "Usabilidad", "P1 - ALTA", "FFF0E0"),
    ("MR-01", "Exportacion a Excel/CSV en todos los modulos", "Mejora", "P1 - ALTA", "FFFDE0"),
    ("MR-02", "Actualizacion de costo al registrar compra", "Mejora", "P2 - MEDIA", "FFFDE0"),
    ("MR-03", "ABM de Proveedores", "Mejora", "P2 - MEDIA", "FFFDE0"),
    ("MR-04", "Vista de Flujo de Caja / Resultado del Periodo", "Mejora", "P2 - MEDIA", "FFFDE0"),
    ("MR-05", "Numeracion correlativa de pedidos", "Mejora", "P2 - MEDIA", "FFFDE0"),
    ("PI-05", "Alerta de stock armable en combos", "Usabilidad", "P2 - MEDIA", "FFF0E0"),
    ("PI-07", "Flujo de entrega parcial / devoluciones", "Usabilidad", "P2 - MEDIA", "FFF0E0"),
    ("RS-03", "Unificar Gastos en Caja & Pagos", "Reorganizacion", "P2 - MEDIA", "E3F0FF"),
    ("MR-06", "Buscador global en el ERP", "Mejora", "P3 - BAJA", "FFFDE0"),
    ("PI-06", "Renombrar 'Compras & Stock'", "Usabilidad", "P3 - BAJA", "FFF0E0"),
    ("PI-08", "Renombrar estado 'Terminado / Listo'", "Usabilidad", "P3 - BAJA", "FFF0E0"),
    ("MR-07", "Corregir texto header Empleados", "Mejora", "P3 - BAJA", "FFFDE0"),
]

for code, desc, cat, prio, bg in priority_data:
    row = priority_table.add_row().cells
    for cell, txt in zip(row, [code, desc, cat, prio]):
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        r = p.add_run(txt)
        r.font.size = Pt(9); r.font.name = "Calibri"
        if txt == code:
            r.bold = True

col_widths = [Cm(2.2), Cm(7.2), Cm(3.5), Cm(3.3)]
for row in priority_table.rows:
    for cell, w in zip(row.cells, col_widths):
        cell.width = w

spacer(2)

# ══════════════════════════════════════════════════════════════════════════════
# CONCLUSION FINAL
# ══════════════════════════════════════════════════════════════════════════════
heading1("Conclusion del Auditor")
body(
    "El ERP de Navidad y Empresas es un sistema funcionalmente bien pensado para el negocio que representa. "
    "El diseno de flujos comerciales (Lead - Cotizacion - Venta - Armado - Entrega) es correcto y refleja la realidad operativa. "
    "Las funcionalidades de combos, logistica diferenciada y liquidacion de personal muestran un profundo entendimiento del negocio."
)
body(
    "Sin embargo, la base tecnica de localStorage y la ausencia de reserva de stock formal son dos limitaciones que no son negociables "
    "para una operacion real de escala. Estas dos fallas deben resolverse antes de incorporar empleados al sistema, "
    "antes de cargar datos reales de clientes, y antes de la proxima temporada navidena."
)
body(
    "La recomendacion profesional es: utilizar la version actual como prototipo funcional para validar flujos y entrenamiento, "
    "mientras se construye en paralelo la version con backend real. Las reorganizaciones de modulos y las mejoras de usabilidad "
    "pueden implementarse progresivamente sin riesgo operativo."
)
body(
    "El sistema tiene una base solida. Con las correcciones indicadas, puede convertirse en un ERP robusto y diferenciado para el rubro."
)

spacer()
cierre = doc.add_paragraph()
cierre.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para_spacing(cierre, before=100, after=0)
r_c = cierre.add_run("Auditoria elaborada - Junio 2026 | ERP Navidad y Empresas")
r_c.italic = True; r_c.font.size = Pt(9); r_c.font.name = "Calibri"
r_c.font.color.rgb = RGBColor(150, 150, 150)

# ─── SAVE ──────────────────────────────────────────────────────────────────────
output_path = "/sessions/charming-wonderful-ritchie/mnt/ERP Navidad y Empresas/Auditoria_ERP_Navidad_y_Empresas.docx"
doc.save(output_path)
print(f"Guardado: {output_path}")
